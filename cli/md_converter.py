#!/usr/bin/env python3
"""Markdown to DOCX converter."""

import io
import mistune
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_run_monospace(run):
    run.font.name = "Courier New"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.insert(0, rFonts)


def _extract_text(nodes):
    """Flatten nodes to plain text."""
    parts = []
    for node in nodes:
        if node.get("type") == "text":
            parts.append(node.get("raw", ""))
        elif node.get("type") == "codespan":
            parts.append(node.get("raw", ""))
        elif node.get("children"):
            parts.append(_extract_text(node["children"]))
        else:
            parts.append(node.get("raw", ""))
    return "".join(parts)


def _render_inline(paragraph, nodes):
    """Add inline nodes as runs to an existing paragraph."""
    if not nodes:
        return
    for node in nodes:
        ntype = node.get("type")
        if ntype == "text":
            paragraph.add_run(node.get("raw", ""))
        elif ntype == "strong":
            run = paragraph.add_run(_extract_text(node.get("children", [])))
            run.bold = True
        elif ntype == "emphasis":
            run = paragraph.add_run(_extract_text(node.get("children", [])))
            run.italic = True
        elif ntype == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            _set_run_monospace(run)
        elif ntype == "link":
            url = node.get("attrs", {}).get("url", "")
            label = _extract_text(node.get("children", []))
            text = f"{label} ({url})" if url and url != label else label
            paragraph.add_run(text)
        elif ntype in ("softlinebreak", "linebreak"):
            paragraph.add_run(" ")
        elif ntype in ("html_inline", "inline_html"):
            pass
        else:
            raw = node.get("raw", "")
            if raw:
                paragraph.add_run(raw)
            elif node.get("children"):
                _render_inline(paragraph, node["children"])


def _render_block(doc, node):
    """Render a block-level AST node into the document."""
    ntype = node.get("type")

    if ntype == "heading":
        level = node.get("attrs", {}).get("level", 1)
        text = _extract_text(node.get("children", []))
        doc.add_heading(text, level=level)

    elif ntype == "paragraph":
        p = doc.add_paragraph()
        _render_inline(p, node.get("children", []))

    elif ntype == "block_code":
        raw = node.get("raw", "")
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(raw)
        _set_run_monospace(run)

    elif ntype == "block_quote":
        for child in node.get("children", []):
            if child.get("type") == "paragraph":
                p = doc.add_paragraph(style="Quote")
                _render_inline(p, child.get("children", []))
            else:
                _render_block(doc, child)

    elif ntype == "list":
        ordered = node.get("attrs", {}).get("ordered", False)
        style = "List Number" if ordered else "List Bullet"
        for item in node.get("children", []):
            p = doc.add_paragraph(style=style)
            for child in item.get("children", []):
                if child.get("type") == "paragraph":
                    _render_inline(p, child.get("children", []))
                elif child.get("type") == "list":
                    _render_block(doc, child)
                else:
                    _render_inline(p, [child])

    elif ntype == "table":
        children = node.get("children", [])
        head_node = next((c for c in children if c.get("type") == "table_head"), None)
        body_node = next((c for c in children if c.get("type") == "table_body"), None)

        # table_head children are table_cell nodes directly
        head_cells = head_node.get("children", []) if head_node else []
        # table_body children are table_row nodes, each with table_cell children
        body_rows = body_node.get("children", []) if body_node else []

        cols = len(head_cells) if head_cells else (
            len(body_rows[0].get("children", [])) if body_rows else 1
        )
        total_rows = (1 if head_cells else 0) + len(body_rows)
        if total_rows == 0 or cols == 0:
            return

        table = doc.add_table(rows=total_rows, cols=cols)
        table.style = "Table Grid"

        row_idx = 0
        if head_cells:
            for col_idx, cell_node in enumerate(head_cells):
                if col_idx < cols:
                    table.cell(row_idx, col_idx).text = _extract_text(
                        cell_node.get("children", [])
                    )
            row_idx += 1

        for row_node in body_rows:
            for col_idx, cell_node in enumerate(row_node.get("children", [])):
                if col_idx < cols:
                    table.cell(row_idx, col_idx).text = _extract_text(
                        cell_node.get("children", [])
                    )
            row_idx += 1

    elif ntype == "thematic_break":
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif ntype in ("blank_line", "html_block"):
        pass

    else:
        for child in node.get("children", []):
            _render_block(doc, child)


def convert_md(filepath):
    """Convert a Markdown file to DOCX bytes."""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()

    md = mistune.create_markdown(renderer=None, plugins=["table"])
    tokens = md(text)

    doc = Document()
    for node in tokens:
        _render_block(doc, node)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
